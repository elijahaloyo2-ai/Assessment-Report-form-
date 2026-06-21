import streamlit as st
import pandas as pd
import io
import zipfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

st.set_page_config(page_title="KEA Report Generator", page_icon="📝", layout="wide")

st.title("📝 KEA COMPREHENSIVE SCHOOL")
st.subheader("Automated Student Assessment Report Generator")

# Helper function to style cells (Background Color)
def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Helper function to add inner cell padding
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report_document(student_data):
    doc = Document()
    
    # Page setup - 0.75 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # 1. SCHOOL HEADER
    p_school = doc.add_paragraph()
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_school = p_school.add_run("KEA COMPREHENSIVE SCHOOL\n")
    run_school.font.name = 'Arial'
    run_school.font.size = Pt(20)
    run_school.font.bold = True
    run_school.font.color.rgb = RGBColor(26, 54, 93) # Deep Navy

    p_address = doc.add_paragraph()
    p_address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_address = p_address.add_run("P.O BOX 557-40404 SUNA MIGORI\n")
    run_address.font.name = 'Arial'
    run_address.font.size = Pt(11)
    run_address.font.italic = True
    run_address.font.color.rgb = RGBColor(74, 85, 104)

    run_sub = p_address.add_run("STUDENT ASSESSMENT REPORT")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(43, 108, 176) # Royal Blue

    # Decorative horizontal line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.add_run("__________________________________________________________________").font.color.rgb = RGBColor(226, 232, 240)

    # 2. STUDENT DETAILS BLOCK
    details_table = doc.add_table(rows=3, cols=4)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details_table.autofit = False
    
    # Smart structural search for Rank/Position field variation
    rank_value = "-"
    for r_key in ['RANK', 'POSITION', 'RA', 'RANK / POSITION']:
        if r_key in student_data and str(student_data[r_key]).strip() != "-":
            rank_value = str(student_data[r_key])
            break

    labels = [
        ("NAME:", str(student_data.get('NAME', ''))),
        ("ADM NO:", str(student_data.get('ADM NO', ''))),
        ("ASSESSMENT NO:", str(student_data.get('ASSESSMENT NUMBER', student_data.get('ASSESSMENT NO', student_data.get('ENT NO', ''))))),
        ("TERM:", str(student_data.get('TERM', ''))),
        ("YEAR:", str(student_data.get('YEAR', ''))),
        ("RANK / POSITION:", rank_value)
    ]
    
    idx = 0
    for r in range(3):
        for c in [0, 2]:
            if idx < len(labels):
                lbl, val = labels[idx]
                # Label cell
                cell_lbl = details_table.cell(r, c)
                cell_lbl.width = Inches(1.5)
                p_lbl = cell_lbl.paragraphs[0]
                p_lbl.add_run(lbl).font.bold = True
                p_lbl.runs[0].font.size = Pt(10)
                p_lbl.runs[0].font.name = 'Arial'
                
                # Value cell
                cell_val = details_table.cell(r, c+1)
                cell_val.width = Inches(2.0)
                p_val = cell_val.paragraphs[0]
                p_val.add_run(val)
                p_val.runs[0].font.size = Pt(10)
                p_val.runs[0].font.name = 'Arial'
                idx += 1

    doc.add_paragraph() # Spacer

    # 3. ACADEMICS PERFORMANCE TABLE
    learning_areas = [
        ("901", "ENGLISH"),
        ("902", "KISWAHILI"),
        ("903", "MATHEMATICS"),
        ("905", "INTEGRATED SCIENCE"),
        ("912", "PRETECHNICAL STUDIES"),
        ("907", "SOCIAL STUDIES"),
        ("908", "CRE"),
        ("909", "AGRICULTURE"),
        ("911", "CREATIVE ARTS AND SPORTS")
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = ["Code", "Learning Area", "Score / Grade", "Performance Level"]
    widths = [Inches(1.0), Inches(2.8), Inches(1.2), Inches(2.0)]
    
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], "1A365D") # Navy header
        set_cell_margins(hdr_cells[i], top=120, bottom=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = 'Arial'

    # Populate table lines with fallback handlers
    for code, area in learning_areas:
        row_cells = table.add_row().cells
        for i in range(4):
            row_cells[i].width = widths[i]
            set_cell_margins(row_cells[i], top=80, bottom=80)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        row_cells[0].text = code
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = area
        
        # --- SCORE / GRADE MAPPING ---
        score_keys = [f"{area} SCORE", f"{area} GRADE", area]
        if area == "CREATIVE ARTS AND SPORTS":
            score_keys.insert(0, "C.A.S")
            score_keys.insert(1, "CREATIVE ARTS AND SPORTS SCORE")
            
        score_val = "-"
        for sk in score_keys:
            if sk in student_data and str(student_data[sk]).strip() != "-":
                score_val = str(student_data[sk])
                break
        row_cells[2].text = score_val
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- PERFORMANCE LEVEL MAPPING ---
        perf_keys = [f"{area} LEVEL", f"{area} PERFORMANCE", f"{area} P.LEVEL", "P.LEVEL", "PERFORMANCE LEVEL"]
        if area == "CREATIVE ARTS AND SPORTS":
            perf_keys.insert(0, "P.LEVEL")
            
        perf_val = "-"
        for pk in perf_keys:
            if pk in student_data and str(student_data[pk]).strip() != "-":
                perf_val = str(student_data[pk])
                break
        row_cells[3].text = perf_val
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Typography formatting rules
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9.5)

    doc.add_paragraph() # Spacer

    # 4. OVERALL SUMMARY BLOCK
    summary_table = doc.add_table(rows=2, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    total_val = student_data.get('TOTAL MARKS', student_data.get('TOTAL SCORE', student_data.get('TOTAL', '-')))
    gen_perf_val = student_data.get('GENERAL PERFORMANCE LEVEL', student_data.get('GENERAL PERFORMANCE', student_data.get('PERFORMANCE LEVEL', '-')))
    
    summary_data = [
        ("TOTAL SCORE:", str(total_val)),
        ("GENERAL PERFORMANCE LEVEL:", str(gen_perf_val))
    ]
    
    for r in range(2):
        lbl, val = summary_data[r]
        c0 = summary_table.cell(r, 0)
        c1 = summary_table.cell(r, 1)
        c0.width = Inches(3.5)
        c1.width = Inches(3.5)
        set_cell_margins(c0, top=100, bottom=100)
        set_cell_margins(c1, top=100, bottom=100)
        set_cell_background(c0, "F7FAFC")
        set_cell_background(c1, "F7FAFC")
        
        p0 = c0.paragraphs[0]
        p0.add_run(lbl).font.bold = True
        p0.runs[0].font.name = 'Arial'
        
        p1 = c1.paragraphs[0]
        p1.add_run(val).font.bold = True
        p1.runs[0].font.color.rgb = RGBColor(43, 108, 176)
        p1.runs[0].font.name = 'Arial'

    doc.add_paragraph()

    # 5. REMARKS AND DATES SECTION
    p_remarks = doc.add_paragraph()
    p_remarks.add_run("Class Teacher's Comment: ").font.bold = True
    
    comment_keys = ['TEACHER COMMENT', 'CLASS TEACHER COMMENT', 'COMMENT', 'COMMENTS', 'REMARKS']
    comment = "-"
    for ck in comment_keys:
        if ck in student_data and str(student_data[ck]).strip() != "-":
            comment = str(student_data[ck])
            break
            
    p_remarks.add_run(f'"{comment}"\n\n')
    p_remarks.runs[0].font.name = 'Arial'
    if len(p_remarks.runs) > 1: p_remarks.runs[1].font.name = 'Arial'

    # Term Dates Row Setup
    date_table = doc.add_table(rows=1, cols=2)
    date_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    c_close = date_table.cell(0, 0)
    c_open = date_table.cell(0, 1)
    c_close.width = Inches(3.5)
    c_open.width = Inches(3.5)
    
    p_close = c_close.paragraphs[0]
    p_close.add_run("Closing Date: ").font.bold = True
    p_close.add_run(str(student_data.get('CLOSING DATE', student_data.get('CLOSING', 'N/A'))))
    
    p_open = c_open.paragraphs[0]
    p_open.add_run("Opening Date: ").font.bold = True
    p_open.add_run(str(student_data.get('OPENING DATE', student_data.get('OPENING', 'N/A'))))
    
    for run in p_close.runs + p_open.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    # 6. SIGN-OFF BLOCK
    doc.add_paragraph("\n\n")
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sign = p_sign.add_run("Signature: ___________________________\n")
    run_sign.font.bold = True
    run_sign.font.name = 'Arial'
    run_school_head = p_sign.add_run("School Principal / Head Teacher       ")
    run_school_head.font.italic = True
    run_school_head.font.size = Pt(9)
    run_school_head.font.name = 'Arial'

    return doc


# --- WEB INTERFACE WORKFLOW ---
uploaded_excel = st.file_uploader("Upload Student Data Excel Sheet (.xlsx)", type=["xlsx"])

if uploaded_excel:
    try:
        # Load data frames and clean spaces, then scale column strings to uniform upper-case
        df = pd.read_excel(uploaded_excel)
        df.columns = df.columns.str.strip().str.upper()
        
        st.success("Excel data loaded successfully!")
        st.write("### Preview of Uploaded Student Data", df.head())
        
        if st.button("Generate Assessment Reports"):
            with st.spinner("Processing files cleanly... Please wait."):
                
                zip_buffer = io.BytesIO()
                
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                    for index, row in df.iterrows():
                        student_data = row.to_dict()
                        
                        # Data scrubbing and cleaning loop
                        cleaned_data = {}
                        for k, v in student_data.items():
                            if isinstance(v, float) and v.is_integer():
                                cleaned_data[k] = str(int(v))
                            elif pd.isna(v) or str(v).strip().lower() == 'none':
                                cleaned_data[k] = "-"
                            else:
                                cleaned_data[k] = str(v)
                        
                        # Generate dynamic word files matching data row context
                        doc = create_report_document(cleaned_data)
                        
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        admin_no = cleaned_data.get('ADM NO', f"index_{index}")
                        name = cleaned_data.get('NAME', 'Student').replace(" ", "_")
                        filename = f"{admin_no}_{name}_Assessment_Report.docx".replace("/", "-")
                        
                        zip_file.writestr(filename, doc_io.getvalue())
                
                zip_buffer.seek(0)
                
                st.success("🎉 All KEA Comprehensive School assessment reports generated successfully!")
                st.download_button(
                    label="📥 Download All Reports (ZIP File)",
                    data=zip_buffer,
                    file_name="KEA_Comprehensive_School_Reports.zip",
                    mime="application/zip"
                )
                
    except Exception as e:
        st.error(f"An error occurred: {e}")
